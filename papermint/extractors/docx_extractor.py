"""Word document extractor module for PaperMint."""

from __future__ import annotations

import io
import logging

import docx

from papermint.errors import CorruptedDocumentError
from papermint.extractors.base import BaseExtractor, ExtractedDocument

logger = logging.getLogger(__name__)

#: Paragraph styles that mark a structural heading in a Word document.
_HEADING_STYLES = ("Heading", "Title", "Subtitle")


class DocxExtractor(BaseExtractor):
    """Extractor for Word (.docx) files using python-docx."""

    name = "Word"
    mime_types = frozenset(
        {"application/vnd.openxmlformats-officedocument.wordprocessingml.document"}
    )
    extensions = frozenset({"docx"})

    def extract_text(self, file_bytes: bytes) -> str:
        """Extract text from a Word document.

        Args:
            file_bytes: The raw .docx bytes.

        Returns:
            The document text with paragraph and table structure preserved.

        Raises:
            CorruptedDocumentError: If the archive cannot be opened.
        """
        return self.extract_document(file_bytes).text

    def extract_document(self, file_bytes: bytes) -> ExtractedDocument:
        """Extract text along with a section count.

        Headings are surrounded by blank lines so that the bibliography
        detector can match a ``References`` heading on a line of its own.

        Args:
            file_bytes: The raw .docx bytes.

        Returns:
            An :class:`ExtractedDocument` describing the file.

        Raises:
            CorruptedDocumentError: If the archive cannot be opened.
        """
        try:
            doc = docx.Document(io.BytesIO(file_bytes))
        except Exception as exc:
            logger.exception("Failed to open DOCX stream")
            raise CorruptedDocumentError(
                "This Word document could not be opened. It may be corrupted, or it "
                "may be an older .doc file rather than .docx.",
                remedy="Open it in Word and save it again as .docx, then re-upload.",
            ) from exc

        blocks: list[str] = []
        seen_table_text: set[str] = set()

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            style_name = getattr(para.style, "name", "") or ""
            if style_name.startswith(_HEADING_STYLES):
                blocks.extend(["", text, ""])
            else:
                blocks.append(text)

        for table in doc.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells if c.text.strip()]
                if not cells:
                    continue
                row_text = " | ".join(dict.fromkeys(cells))
                if row_text not in seen_table_text:
                    seen_table_text.add(row_text)
                    blocks.append(row_text)

        try:
            section_count = len(doc.sections)
        except Exception:  # pragma: no cover - malformed but readable documents
            section_count = 0

        return ExtractedDocument(text="\n".join(blocks), page_count=section_count)
