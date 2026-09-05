"""The reference formatter: set what was parsed, and explain the style it is set in.

Everywhere else in PaperMint reads documents. This page writes: it takes the
references the analyzer produced, or a single reference pasted by hand, and
renders them as a finished reference list in APA, MLA, IEEE or Chicago, beside
an account of what that style is and why it is built the way it is.

It is deliberately the one page that consumes another page's work. The
analyzer's citations live in session state, so they are already here when the
reader arrives, which is also the plainest demonstration that moving between
pages costs nothing.

The honesty principle applies to formatting as much as to parsing. Every
rendered entry names the elements the source never supplied, and no title is
recased on the reader's behalf, because deciding which words are proper nouns
is exactly the judgement a machine gets wrong.
"""

from __future__ import annotations

import io
import logging

import streamlit as st

from papermint.errors import PaperMintError
from papermint.formatters.reference_formatter import (
    FormattedReference,
    StyleGuide,
    format_reference_list,
    formattable_styles,
    style_guide,
)
from papermint.models import Citation
from papermint.pipeline import PipelineService
from papermint.ui.components.citation_card import render_citation_card
from papermint.ui.components.primitives import (
    chip_row,
    definition_list,
    empty_state,
    notice,
    page_header,
    prose,
    section_header,
)
from papermint.ui.html import compact, esc
from papermint.ui.html import render as render_html
from papermint.ui.state import restore, retain

logger = logging.getLogger(__name__)

#: Where the analyzer leaves its work, including any reader edits.
_ANALYZER_CITATIONS = "pm_extract_citations"
_ANALYZER_RESULT = "pm_extract_result"

_STYLE_KEY = "pm_style_choice"
_SOURCE_KEY = "pm_style_source"
_PASTE_KEY = "pm_style_paste"
_PARSED_KEY = "pm_style_parsed"
_COMPARE_KEY = "pm_style_compare"

_STICKY_KEYS = (_STYLE_KEY, _SOURCE_KEY, _PASTE_KEY)

#: The two ways a reference can arrive on this page.
_FROM_ANALYZER = "The references I just extracted"
_FROM_PASTE = "One reference I paste"

#: How far into a list the entrance cascade keeps growing before the delay is
#: held constant, so a hundred-entry bibliography does not take six seconds to
#: finish arriving.
_MAX_REVEAL_STEP = 10


def _guides() -> dict[str, StyleGuide]:
    """Return every style guide, keyed by the name shown on the selector.

    Returns:
        A mapping of short style name to guide, in presentation order.
    """
    return {style_guide(style).short_name: style_guide(style) for style in formattable_styles()}


# ---------------------------------------------------------------------------
# The guide
# ---------------------------------------------------------------------------


def _render_chain(guide: StyleGuide) -> None:
    """Render a style's elements as a numbered chain, in order.

    Seeing the order and the punctuation together is what turns a style from an
    arbitrary set of rules into a shape a writer can reproduce, which is why
    this is drawn rather than listed.

    Args:
        guide: The style being explained.
    """
    items: list[str] = []
    for position, (name, rule) in enumerate(guide.elements):
        items.append(
            f'<div class="pm-chain-item" style="--pm-step:{min(position, _MAX_REVEAL_STEP)};">'
            f'<div class="pm-chain-num">{position + 1:02d}</div>'
            "<div>"
            f'<div class="pm-chain-name">{esc(name)}</div>'
            f'<div class="pm-chain-rule">{esc(rule)}</div>'
            "</div>"
            "</div>"
        )
    render_html(compact(f'<div class="pm-chain">{"".join(items)}</div>'))


def _render_guide(guide: StyleGuide) -> None:
    """Render everything the interface knows about one style.

    Args:
        guide: The style being explained.
    """
    section_header(guide.name, guide.disciplines)
    prose(guide.principle)
    st.write("")
    chip_row(
        [
            ("library", f"List heading: {guide.list_heading}"),
            ("quote", f"In text: {guide.in_text.split('.')[0]}"),
        ],
        accent_first=True,
    )

    section_header("The elements, in order")
    _render_chain(guide)

    section_header("What sets it apart")
    definition_list(
        [
            ("Ordering", guide.ordering),
            ("Citing in text", guide.in_text),
            *[(f"Detail {n}", detail) for n, detail in enumerate(guide.distinctives, start=1)],
        ]
    )

    section_header("A finished entry")
    render_html(
        compact(f'<div class="pm-reflist"><div class="pm-refline">{esc(guide.sample)}</div></div>')
    )
    st.write("")
    notice("One thing PaperMint will not do for you", guide.caveat, tone="info")


# ---------------------------------------------------------------------------
# The formatted output
# ---------------------------------------------------------------------------


def _reference_markup(rendered: FormattedReference, position: int) -> str:
    """Render one formatted reference as a hanging-indent line.

    Args:
        rendered: The formatted reference.
        position: Its place in the entrance cascade.

    Returns:
        The line's HTML.
    """
    text = esc(rendered.text)
    if rendered.italic:
        # Both strings pass through the same escaping, so the substring still
        # matches. Only the first occurrence is emphasised: a journal whose
        # name also appears in the title should be italic once, not twice.
        needle = esc(rendered.italic)
        text = text.replace(needle, f"<em>{needle}</em>", 1)

    marker = f'<span class="pm-refmark">{esc(rendered.marker)}</span>' if rendered.marker else ""
    gap = ""
    if rendered.omitted:
        gap = (
            '<span class="pm-refgap">Absent from the source, so absent here: '
            f"{esc(', '.join(rendered.omitted))}</span>"
        )
    step = min(position, _MAX_REVEAL_STEP)
    return f'<div class="pm-refline" style="--pm-step:{step};">{marker}{text}{gap}</div>'


def _render_reference_list(rendered: list[FormattedReference]) -> None:
    """Render a whole formatted list the way it is set on paper.

    Args:
        rendered: The formatted references, in list order.
    """
    lines = "".join(_reference_markup(ref, position) for position, ref in enumerate(rendered))
    render_html(compact(f'<div class="pm-reflist">{lines}</div>'))


def _plain_text(rendered: list[FormattedReference], heading: str) -> str:
    """Assemble the list as plain text, ready to paste into a manuscript.

    Args:
        rendered: The formatted references.
        heading: The style's own name for the list.

    Returns:
        The list as text.
    """
    body = "\n\n".join(f"{ref.marker} {ref.text}".strip() for ref in rendered)
    return f"{heading}\n\n{body}\n"


def _docx_bytes(rendered: list[FormattedReference], guide: StyleGuide) -> io.BytesIO:
    """Render the formatted reference list into a Word document.

    Applies standard academic hanging indents (0.5 inch left indent, -0.5 inch
    first line indent) and preserves italics for periodical and book titles.

    Args:
        rendered: The formatted references.
        guide: The chosen style guide.

    Returns:
        A BytesIO stream containing the .docx file.
    """
    import docx
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches, Pt

    doc = docx.Document()

    heading = doc.add_heading(guide.list_heading, level=0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph(
        f"Generated by PaperMint · {len(rendered)} references formatted in {guide.short_name}"
    )
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.style = "Subtitle"

    doc.add_paragraph()

    for ref in rendered:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.5)
        p.paragraph_format.first_line_indent = Inches(-0.5)
        p.paragraph_format.space_after = Pt(6)

        if ref.marker:
            run_m = p.add_run(f"{ref.marker} ")
            run_m.bold = True

        if ref.italic and ref.italic in ref.text:
            before, needle, after = ref.text.partition(ref.italic)
            if before:
                p.add_run(before)
            run_it = p.add_run(needle)
            run_it.italic = True
            if after:
                p.add_run(after)
        else:
            p.add_run(ref.text)

    output = io.BytesIO()
    doc.save(output)
    output.seek(0)
    return output


def _render_output(citations: list[Citation], guide: StyleGuide) -> None:
    """Render the formatted reference list and the ways to take it away.

    Args:
        citations: The citations to render.
        guide: The chosen style.
    """
    rendered = format_reference_list(citations, guide.style)
    incomplete = sum(1 for ref in rendered if not ref.complete)

    section_header(
        guide.list_heading,
        f"{len(rendered)} entries · {guide.ordering.split('.')[0].lower()}",
    )
    if incomplete:
        notice(
            f"{incomplete} of {len(rendered)} entries are missing an element",
            "Those elements were never found in the source document, so they are "
            "left out here rather than invented. Each affected entry names what it "
            "is missing, and the analyzer's inline editor can supply them.",
            tone="caution",
        )
        st.write("")

    _render_reference_list(rendered)

    st.write("")
    text = _plain_text(rendered, guide.list_heading)
    docx_stream = _docx_bytes(rendered, guide)
    copy_col, txt_col, docx_col = st.columns([2, 1, 1])
    with copy_col.popover("Copy as plain text", use_container_width=True):
        st.caption("Use the copy button in the corner of the block.")
        st.code(text, language=None)
    txt_col.download_button(
        "Download TXT",
        data=text.encode("utf-8"),
        file_name=f"{guide.short_name.lower()}_{guide.list_heading.lower().replace(' ', '_')}.txt",
        mime="text/plain",
        use_container_width=True,
        key="pm_style_download",
    )
    docx_col.download_button(
        "Download Word (.docx)",
        data=docx_stream.getvalue(),
        file_name=f"{guide.short_name.lower()}_{guide.list_heading.lower().replace(' ', '_')}.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        use_container_width=True,
        key="pm_style_docx_download",
    )


def _render_comparison(citations: list[Citation]) -> None:
    """Show one reference in all four styles at once.

    A style is easiest to learn against another one. The same entry set four
    ways makes the differences visible in a single glance: names that invert, a
    year that moves to the end, quotation marks that come and go.

    Args:
        citations: The citations available to compare.
    """
    section_header("The same reference, four ways")
    labels = [f"{position + 1:02d}. {c.display_title}" for position, c in enumerate(citations)]
    chosen = st.selectbox(
        "Reference to compare",
        options=range(len(citations)),
        format_func=lambda index: labels[index],
        key=_COMPARE_KEY,
        label_visibility="collapsed",
    )

    citation = citations[chosen]
    for position, style in enumerate(formattable_styles()):
        guide = style_guide(style)
        rendered = format_reference_list([citation], style)[0]
        chip_row([("quote", guide.short_name)], accent_first=True)
        render_html(
            compact(f'<div class="pm-reflist">{_reference_markup(rendered, position)}</div>')
        )
        st.write("")


# ---------------------------------------------------------------------------
# Sources
# ---------------------------------------------------------------------------


def _parse_pasted() -> Citation | None:
    """Read a pasted reference and parse it through the service.

    Returns:
        The parsed citation, or None when nothing has been submitted yet.
    """
    st.text_area(
        "Reference",
        placeholder=(
            "Ambler, Marjane. Women Leaders in Indian Education. "
            "Tribal College, vol. 3, no. 4, 1992, pp. 10-15."
        ),
        height=110,
        key=_PASTE_KEY,
        label_visibility="collapsed",
    )
    if not st.button("Read this reference", type="primary", key="pm_style_parse"):
        return st.session_state.get(_PARSED_KEY)

    try:
        parsed = PipelineService().parse_reference(st.session_state.get(_PASTE_KEY, ""))
    except PaperMintError as err:
        notice(
            "That could not be read",
            str(err),
            tone="caution",
            details=[err.remedy] if err.remedy else None,
        )
        return None
    except Exception:
        logger.exception("Unexpected failure parsing a pasted reference")
        notice(
            "Something went wrong",
            "An unexpected error interrupted parsing. The details were written to "
            "the application log.",
            tone="critical",
        )
        return None

    st.session_state[_PARSED_KEY] = parsed
    return parsed


def _analyzer_citations() -> list[Citation]:
    """Return the citations the analyzer page left in session state.

    Returns:
        The citations, or an empty list when nothing has been analysed.
    """
    citations = st.session_state.get(_ANALYZER_CITATIONS) or []
    return [c for c in citations if isinstance(c, Citation)]


# ---------------------------------------------------------------------------
# Page
# ---------------------------------------------------------------------------


def render() -> None:
    """Render the reference formatter page."""
    restore(*_STICKY_KEYS)
    guides = _guides()
    st.session_state.setdefault(_STYLE_KEY, next(iter(guides)))

    page_header(
        "Reference formatter",
        "Set your references in APA, MLA, IEEE or Chicago, and see what each "
        "style actually asks for. Nothing is invented: an element the source "
        "never supplied is left out and named.",
        eyebrow="Format",
        eyebrow_icon="quote",
    )

    available = _analyzer_citations()
    options = [_FROM_ANALYZER, _FROM_PASTE] if available else [_FROM_PASTE]
    if st.session_state.get(_SOURCE_KEY) not in options:
        st.session_state[_SOURCE_KEY] = options[0]

    source_col, style_col = st.columns([2, 3])
    with source_col:
        source = st.radio("Where the references come from", options=options, key=_SOURCE_KEY)
    with style_col:
        chosen = st.radio(
            "Style",
            options=list(guides),
            horizontal=True,
            key=_STYLE_KEY,
            help="Every reference below is rendered in the style selected here.",
        )
    guide = guides[chosen]

    citations: list[Citation] = []
    if source == _FROM_ANALYZER:
        result = st.session_state.get(_ANALYZER_RESULT)
        origin = getattr(result, "source_filename", "")
        st.caption(
            f"{len(available)} references"
            + (f" from {origin}" if origin else "")
            + ", including any corrections you made on the analyzer page."
        )
        citations = available
    else:
        parsed = _parse_pasted()
        if parsed is not None:
            st.write("")
            render_citation_card(parsed, 1)
            citations = [parsed]

    st.divider()

    if citations:
        _render_output(citations, guide)
        st.divider()
        _render_comparison(citations)
        st.divider()
    elif source == _FROM_ANALYZER:
        empty_state(
            "Nothing to format yet",
            "Analyse a document first and its references will be waiting here.",
            icon_name="library",
        )
    else:
        empty_state(
            "Paste a reference above",
            "One line from a Works Cited page is enough. PaperMint reads it into "
            "fields and sets it in all four styles.",
            icon_name="quote",
        )

    _render_guide(guide)
    retain(*_STICKY_KEYS)


__all__ = ["render"]
