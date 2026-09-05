"""Tests for the presentation layer.

The component tests assert on generated markup, which is where the interface's
two historical failure modes lived: unescaped document text corrupting a card,
and indented HTML being rendered as a Markdown code block. The page tests run
the real application through Streamlit's own harness and assert that no screen
raises.
"""

from __future__ import annotations

import io
from pathlib import Path

import pytest
from streamlit.testing.v1 import AppTest

from papermint.models import (
    Author,
    BatchFileResult,
    BatchResult,
    Citation,
    CitationStyle,
    ConfidenceBand,
    ExtractionResult,
)
from papermint.pipeline import PipelineStage
from papermint.ui.components.citation_browser import _matches, _sorted_citations
from papermint.ui.components.citation_card import (
    _card_markup,
    _parse_author_field,
    citation_preview_text,
)
from papermint.ui.components.export_panel import EXPORT_FORMATS, safe_filename
from papermint.ui.components.file_uploader import read_upload, upload_signature
from papermint.ui.components.progress import _flow_markup
from papermint.ui.html import clamp, compact, dot_join, esc
from papermint.ui.icons import available_icons, icon
from papermint.ui.navigation import route_names
from papermint.ui.pages.batch import _outcome
from papermint.ui.styles import build_stylesheet
from papermint.ui.theme import COLOR, band_color, css_variables


@pytest.fixture
def citation() -> Citation:
    """Return a fully populated citation."""
    return Citation(
        title="Machine learning in citation parsing",
        authors=[Author(given="J. A.", family="Smith")],
        year="2020",
        journal="Journal of Bibliometrics",
        volume="15",
        issue="2",
        pages="103-115",
        doi="10.1016/j.jbi.2020.01.002",
        style=CitationStyle.APA,
        confidence=0.9,
        raw_text="Smith, J. A. (2020). Machine learning in citation parsing.",
    )


# --- HTML helpers ----------------------------------------------------------


def test_escaping_neutralises_markup():
    assert esc('<script>alert("x")</script>') == (
        "&lt;script&gt;alert(&quot;x&quot;)&lt;/script&gt;"
    )


def test_escaping_handles_none():
    assert esc(None) == ""


def test_compact_removes_indentation_that_markdown_would_eat():
    markup = "    <div>\n        <span>hi</span>\n    </div>"
    result = compact(markup)
    assert not result.startswith(" ")
    assert "\n    " not in result
    assert result == "<div><span>hi</span></div>"


def test_clamp_breaks_on_a_word_boundary():
    assert clamp("alpha beta gamma delta", 12).endswith("…")
    assert clamp("short", 40) == "short"


def test_dot_join_drops_empty_parts():
    assert dot_join("a", "", "b") == "a · b"


# --- Theme and icons -------------------------------------------------------


def test_every_token_becomes_a_css_variable():
    variables = css_variables()
    assert "--pm-color-accent: #34D399;" in variables
    assert "--pm-space-4: 16px;" in variables


def test_confidence_bands_map_to_distinct_colours():
    colours = {band_color(band) for band in ConfidenceBand}
    assert len(colours) == 3
    assert band_color(ConfidenceBand.HIGH) == COLOR["positive"]


def test_icons_inherit_the_text_colour():
    svg = icon("leaf")
    assert 'stroke="currentColor"' in svg
    assert 'aria-hidden="true"' in svg


def test_an_unknown_icon_renders_nothing_rather_than_raising():
    assert icon("no-such-icon") == ""


def test_the_icon_set_is_not_empty():
    assert len(available_icons()) >= 15


def test_the_stylesheet_uses_tokens_not_literals():
    css = build_stylesheet()
    assert "--pm-color-accent" in css
    assert "stTabs" in css
    assert "pm-card" in css


# --- Citation card ---------------------------------------------------------


def test_card_shows_every_parsed_field(citation):
    markup = _card_markup(citation, 1)
    assert "Machine learning in citation parsing" in markup
    assert "Smith, J. A." in markup
    assert "Journal of Bibliometrics" in markup
    assert ">15</dd>" in markup
    assert ">2</dd>" in markup
    assert ">103-115</dd>" in markup
    assert "10.1016/j.jbi.2020.01.002" in markup


def test_card_labels_every_field_it_shows(citation):
    markup = _card_markup(citation, 1)
    for label in ("Title", "Authors", "Year", "Journal", "Volume", "Issue", "Pages", "DOI", "Type"):
        assert f">{label}</dt>" in markup, f"the {label} row lost its label"


def test_card_keeps_a_labelled_row_for_an_absent_identity_field():
    # Title, authors and year are the identity of a reference. A card that
    # dropped an absent one would leave the reader to notice the gap.
    sparse = Citation(title="Only a title", confidence=0.2, raw_text="x")
    markup = _card_markup(sparse, 1)
    assert ">Authors</dt>" in markup
    assert ">Year</dt>" in markup
    assert markup.count("pm-field-absent") == 2


def test_card_carries_its_coverage_as_a_meter(citation):
    assert "--pm-meter:90%" in _card_markup(citation, 1)


def test_cards_stagger_their_entrance():
    first = _card_markup(Citation(raw_text="a"), 1, reveal=0)
    later = _card_markup(Citation(raw_text="b"), 2, reveal=3)
    assert "--pm-step:0" in first
    assert "--pm-step:3" in later


def test_card_escapes_document_text():
    hostile = Citation(title="A <b>bold</b> claim & more", raw_text="x", confidence=0.5)
    markup = _card_markup(hostile, 1)
    assert "<b>bold</b>" not in markup
    assert "&lt;b&gt;bold&lt;/b&gt;" in markup


def test_card_never_indents_enough_to_become_a_code_block(citation):
    assert "\n    " not in _card_markup(citation, 1)


def test_card_uses_raw_text_when_no_title_was_parsed():
    unparsed = Citation(raw_text="Some unparseable entry text here", confidence=0.1)
    markup = _card_markup(unparsed, 3)
    assert "Untitled" not in markup
    assert "Some unparseable entry text here" in markup
    assert "is-unparsed" in markup


def test_card_lists_missing_fields_on_low_confidence_entries():
    sparse = Citation(title="Only a title", confidence=0.2, raw_text="x")
    assert "Not found:" in _card_markup(sparse, 1)


def test_card_hides_missing_fields_on_complete_entries(citation):
    assert "Not found:" not in _card_markup(citation, 1)


def test_edited_authors_round_trip_in_either_order():
    authors = _parse_author_field("Smith, J. A.; Robert B. Doe")
    assert [a.family for a in authors] == ["Smith", "Doe"]
    assert authors[1].given == "Robert B."


def test_preview_text_is_one_line(citation):
    assert "\n" not in citation_preview_text(citation)


# --- Export panel ----------------------------------------------------------


def test_filenames_are_made_safe():
    assert safe_filename("My Paper (final)/v2") == "My_Paper_final_v2"
    assert safe_filename("") == "citations"
    assert safe_filename("///") == "citations"


def test_every_export_format_is_serialisable(citation):
    for fmt in EXPORT_FORMATS:
        payload = fmt.serialise([citation])
        assert payload is not None
        if isinstance(payload, str):
            assert payload.strip()


# --- Upload handling -------------------------------------------------------


def test_upload_is_readable_more_than_once():
    buffer = io.BytesIO(b"payload")
    assert read_upload(buffer) == b"payload"
    assert read_upload(buffer) == b"payload"


def test_upload_signature_changes_with_the_options():
    buffer = io.BytesIO(b"payload")
    assert upload_signature(buffer, False) != upload_signature(buffer, True)


# --- Pages -----------------------------------------------------------------


APP_PATH = Path(__file__).resolve().parent.parent / "app.py"


def test_the_app_starts_without_error():
    harness = AppTest.from_file(str(APP_PATH), default_timeout=120)
    harness.run()
    assert not harness.exception, [str(e.value) for e in harness.exception]


@pytest.mark.parametrize("route", route_names())
def test_every_page_renders_without_error(route: str):
    # A page object can only be executed through a navigation, so each route
    # is driven by a navigation restricted to that one page.
    script = (
        "from papermint.ui.styles import inject_custom_css\n"
        "from papermint.ui.navigation import build_navigation\n"
        "inject_custom_css()\n"
        f"build_navigation(only={route!r}).run()\n"
    )
    harness = AppTest.from_string(script, default_timeout=120)
    harness.run()
    assert not harness.exception, [str(e.value) for e in harness.exception]


# --- Processing flow -------------------------------------------------------


def test_the_flow_separates_done_active_and_waiting_stages():
    markup = _flow_markup(PipelineStage.PARSE, previous=25.0)
    assert "is-done" in markup
    assert "is-active" in markup
    assert "is-waiting" in markup


def test_the_flow_animates_on_from_where_it_had_reached():
    # Streamlit remounts the node on every write, so the rail has to be told
    # where it was or the fill would restart from zero at every stage.
    markup = _flow_markup(PipelineStage.PARSE, previous=37.5)
    assert "--pm-from:37.5%" in markup
    assert "--pm-to:62.5%" in markup


def test_the_running_flow_says_what_the_stage_is_doing():
    markup = _flow_markup(PipelineStage.CHARACTERIZE)
    assert "Locating bibliography" in markup
    assert "collecting every reference block" in markup


def test_a_finished_flow_is_still_and_silent():
    markup = _flow_markup(PipelineStage.DONE, animated=False)
    assert "pm-flow-status" not in markup
    assert "is-live" not in markup
    assert markup.count("is-done") == 4


# --- State that survives a page switch --------------------------------------

#: A page that draws one widget, restoring and retaining it around the draw.
_STICKY_SCRIPT = """
import streamlit as st
from papermint.ui.state import restore, retain

restore("pm_demo")
st.session_state.setdefault("pm_demo", "typed")
retain("pm_demo")
"""

#: A page that only restores a bounded integer, as the page slider does.
_CLAMP_SCRIPT = """
from papermint.ui.state import restore_within

restore_within("pm_page", 1, 3)
"""


def test_a_widget_value_outlives_the_loss_of_its_key():
    # Streamlit collects the state of any widget it did not draw on a run,
    # which is what emptied the analyzer when the reader visited another page.
    harness = AppTest.from_string(_STICKY_SCRIPT, default_timeout=60)
    harness.run()
    assert harness.session_state["_pm_kept_pm_demo"] == "typed"

    del harness.session_state["pm_demo"]
    harness.run()
    assert harness.session_state["pm_demo"] == "typed"


def test_a_remembered_page_number_is_clamped_into_the_new_range():
    harness = AppTest.from_string(_CLAMP_SCRIPT, default_timeout=60)
    harness.session_state["_pm_kept_pm_page"] = 9
    harness.run()
    assert harness.session_state["pm_page"] == 3


def test_the_navigation_exposes_every_route():
    assert set(route_names()) == {"home", "extract", "batch", "styles", "about"}


def test_navigation_uses_reference_formatter_title():
    from papermint.ui.navigation import _ROUTES, build_pages

    pages = build_pages()
    assert pages["styles"].title == "Reference formatter"
    styles_route = next(r for r in _ROUTES if r[0] == "styles")
    assert styles_route[1] == "Reference formatter"
    assert styles_route[3] == "reference-formatter"


def test_a_failed_flow_shows_failure_mark_and_retains_completed_stages():
    markup = _flow_markup(
        PipelineStage.PARSE,
        previous=25.0,
        animated=False,
        failed_stage=PipelineStage.PARSE,
        failure_message="Malformed citation block",
    )
    assert "is-failed" in markup
    assert "is-done" in markup
    assert "is-waiting" in markup
    assert "Malformed citation block" in markup
    assert 'class="pm-flow-beacon is-failed"' in markup


def test_reference_formatter_docx_export(citation):
    import docx

    from papermint.formatters.reference_formatter import format_reference_list, style_guide
    from papermint.ui.pages.style_studio import _docx_bytes

    guide = style_guide(CitationStyle.APA)
    rendered = format_reference_list([citation], guide.style)
    stream = _docx_bytes(rendered, guide)
    assert stream.getbuffer().nbytes > 0

    doc = docx.Document(stream)
    assert len(doc.paragraphs) >= 2
    ref_p = doc.paragraphs[-1]
    assert ref_p.paragraph_format.left_indent.inches == 0.5
    assert ref_p.paragraph_format.first_line_indent.inches == -0.5


# --- The batch workbench ---------------------------------------------------

#: Drives the batch page alone, as the route tests do.
_BATCH_SCRIPT = (
    "from papermint.ui.styles import inject_custom_css\n"
    "from papermint.ui.navigation import build_navigation\n"
    "inject_custom_css()\n"
    "build_navigation(only='batch').run()\n"
)


def _batch_run() -> BatchResult:
    """Return a three-file run: one read, one without references, one failed."""
    parsed = Citation(
        title="Machine learning in citation parsing",
        authors=[Author(given="J. A.", family="Smith")],
        year="2020",
        journal="Journal of Bibliometrics",
        confidence=0.9,
        source_file="paper.pdf",
    )
    return BatchResult(
        files=[
            BatchFileResult(
                filename="paper.pdf",
                result=ExtractionResult(
                    citations=[parsed],
                    source_filename="paper.pdf",
                    raw_text="Body text.\nReferences\nSmith, J. A. (2020).",
                ),
            ),
            BatchFileResult(
                filename="poster.png",
                result=ExtractionResult(source_filename="poster.png", raw_text="A poster."),
            ),
            BatchFileResult(filename="broken.pdf", error="The file could not be opened."),
        ],
        duration_ms=1200,
    )


def _batch_harness(**state: object) -> AppTest:
    """Render the batch page over a cached run, as a returning reader sees it."""
    harness = AppTest.from_string(_BATCH_SCRIPT, default_timeout=120)
    harness.session_state["pm_batch_result"] = _batch_run()
    for key, value in state.items():
        harness.session_state[key] = value
    harness.run()
    assert not harness.exception, [str(e.value) for e in harness.exception]
    return harness


def _doc_heads(harness: AppTest) -> list[str]:
    """Return the pane headings on screen. There should never be more than one."""
    return [m.value for m in harness.markdown if '<div class="pm-doc-head">' in m.value]


def test_the_batch_run_shows_one_document_at_a_time():
    # The old page stacked an expander per file and rendered every citation in
    # all of them at once. One pane replaces the stack.
    harness = _batch_harness()
    assert [tab.label for tab in harness.tabs] == ["Documents", "Merged export"]
    assert not [e for e in harness.expander if "paper.pdf" in e.label]

    heads = _doc_heads(harness)
    assert len(heads) == 1
    assert "paper.pdf" in heads[0]


def test_choosing_a_document_puts_it_in_the_pane():
    harness = _batch_harness()
    harness.button(key="pm_batch_pick_1").click().run()

    assert harness.session_state["pm_batch_file"] == 1
    heads = _doc_heads(harness)
    assert len(heads) == 1
    assert "poster.png" in heads[0]


def test_the_merged_export_is_not_below_the_documents():
    # The complaint that started this: the export sat under every file, so
    # reaching it meant scrolling past all of them. It now has its own tab,
    # and each document carries its own export beside its heading.
    harness = _batch_harness()
    assert [d.key for d in harness.tabs[1].download_button] == ["batch_btn"]
    assert [d.key for d in harness.tabs[0].download_button] == ["pm_batch_doc_export_btn"]


def test_a_file_that_failed_is_explained_in_its_own_pane():
    harness = _batch_harness()
    harness.button(key="pm_batch_pick_2").click().run()

    assert harness.session_state["pm_batch_file"] == 2
    assert any("could not be opened" in m.value for m in harness.markdown)


def test_a_selection_left_over_from_a_longer_run_is_pulled_back_into_range():
    harness = _batch_harness(pm_batch_file=9)
    assert harness.session_state["pm_batch_file"] == 0
    assert "paper.pdf" in _doc_heads(harness)[0]


def test_the_switcher_says_how_each_file_turned_out():
    files = _batch_run().files
    assert _outcome(files[0])[1] == "1 reference · 90% read"
    assert _outcome(files[1])[1] == "No bibliography"
    assert _outcome(files[2])[1:] == ("Could not be read", "critical")


def test_a_merged_listing_names_the_file_each_entry_came_from(citation):
    entry = citation.model_copy(update={"source_file": "thesis.pdf"})
    assert "thesis.pdf" in _card_markup(entry, 1, show_source=True)
    # A single document's own list says nothing the reader already knows.
    assert "thesis.pdf" not in _card_markup(entry, 1)


# --- The citation browser --------------------------------------------------


def test_the_browser_leaves_document_order_alone_by_default(citation):
    second = citation.model_copy(update={"title": "A later entry", "year": "1999"})
    rows = _sorted_citations([citation, second], "Document order")
    assert [index for index, _ in rows] == [0, 1]


def test_the_browser_keeps_each_entrys_place_when_it_reorders(citation):
    second = citation.model_copy(update={"title": "A later entry", "year": "1999"})
    rows = _sorted_citations([citation, second], "Year, oldest first")
    # The pair's first member is the position in the unsorted list, which is
    # what an edit has to be written back to.
    assert [index for index, _ in rows] == [1, 0]


def test_the_browser_finds_an_entry_by_the_file_it_came_from(citation):
    entry = citation.model_copy(update={"source_file": "ERIC_ED060699.pdf"})
    assert _matches(entry, "eric_ed060699")
    assert not _matches(entry, "somewhere else")
